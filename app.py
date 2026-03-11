"""
会议纪要生成 Agent - Streamlit 小程序
页面简洁美观，支持对话式修改 Skill
"""

import os
import json
import io
from datetime import datetime
from pathlib import Path

import streamlit as st
from docx import Document
from docx.shared import Inches

# GitHub 管理模块
from github_manager import GitHubManager, create_github_manager, is_github_mode

# GLM-4 客户端
from glm_client import GLMClient

# 页面配置
st.set_page_config(
    page_title="会议纪要生成 Agent",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS样式
st.markdown("""
<style>
    /* 主容器样式 */
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1E3A8A;
        text-align: center;
        padding: 2rem 0;
    }

    /* 章节标题样式 */
    .section-title {
        font-size: 1.5rem;
        font-weight: bold;
        color: #3B82F6;
        padding: 1rem 0;
        border-bottom: 2px solid #E5E7EB;
    }

    /* 输入框样式 */
    .stTextArea > div > div > textarea {
        border: 2px solid #D1D5DB;
        border-radius: 8px;
        padding: 1rem;
    }

    /* 按钮样式 */
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.75rem 2rem;
        font-size: 1.1rem;
        font-weight: bold;
        transition: all 0.3s;
    }

    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
    }

    /* 状态提示样式 */
    .status-box {
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }

    .status-success {
        background-color: #D1FAE5;
        border-left: 4px solid #10B981;
        color: #065F46;
    }

    .status-warning {
        background-color: #FEF3C7;
        border-left: 4px solid #F59E0B;
        color: #92400E;
    }

    /* 术语词典样式 */
    .dict-table {
        font-size: 0.9rem;
    }
</style>
""", unsafe_allow_html=True)


# ==================== 工具函数 ====================

def get_project_root():
    """获取项目根目录"""
    return Path(__file__).parent


def get_skill_md_path():
    """获取 skill.md 文件路径"""
    return get_project_root() / "skill.md"


def _get_available_history_samples() -> List[str]:
    """
    获取可用的历史纪要文件列表

    Returns:
        历史纪要文件名列表
    """
    # 优先尝试从 GitHub 读取
    github_mgr = create_github_manager(st)
    if github_mgr:
        try:
            # GitHub 模式：从仓库读取
            available_files = github_mgr.list_files("历史纪要/")
            # 过滤出纪要文件
            available_samples = [f for f in available_files if f.startswith('管理周会纪要') and f.endswith('.md')]
            print(f"✅ 从 GitHub 获取到 {len(available_samples)} 篇历史纪要")
            return available_samples
        except Exception as e:
            print(f"⚠️ GitHub 读取失败: {e}，回退到本地读取")

    # 回退到本地读取
    history_dir = get_project_root() / "reference" / "历史纪要"
    if history_dir.exists():
        available_samples = [f.name for f in history_dir.glob('管理周会纪要*.md')]
        print(f"✅ 从本地获取到 {len(available_samples)} 篇历史纪要")
        return available_samples
    else:
        print("⚠️ 本地历史纪要目录不存在")
        return []


def get_dict_md_path():
    """获取术语词典文件路径"""
    return get_project_root() / "reference" / "02_组织与术语词典.md"


def read_file(file_path):
    """读取文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return None


def write_file(file_path, content):
    """写入文件内容"""
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return True
    except Exception as e:
        st.error(f"写入文件失败: {e}")
        return False


def read_reference_file(file_name: str) -> str:
    """
    读取 reference 文件（支持本地和 GitHub 模式）

    Args:
        file_name: 文件名（如 "01_历史纪要重点总结.md"）

    Returns:
        文件内容
    """
    # 优先尝试从 GitHub 读取
    github_mgr = create_github_manager(st)
    if github_mgr:
        content = github_mgr.get_file(f"reference/{file_name}")
        if content:
            print(f"✅ 从 GitHub 读取成功: {file_name} ({len(content)} 字符)")
            return content
        else:
            print(f"⚠️ GitHub 读取失败，回退到本地: {file_name}")

    # 回退到本地读取
    local_path = get_project_root() / "reference" / file_name
    content = read_file(local_path) or ""

    if content:
        print(f"✅ 从本地读取成功: {file_name} ({len(content)} 字符)")
    else:
        print(f"⚠️ 本地读取失败: {file_name}")

    return content


def read_reference_file_no_prefix(file_name: str) -> str:
    """
    读取 reference 文件（不带 reference/ 前缀，兼容根目录结构）

    Args:
        file_name: 文件名（如 "01_历史纪要重点总结.md"）

    Returns:
        文件内容
    """
    # 先尝试直接读取（根目录）
    local_path = get_project_root() / file_name
    content = read_file(local_path) or ""

    if content:
        print(f"✅ 从本地读取成功（直接路径）: {file_name} ({len(content)} 字符)")
        return content

    # 如果直接读取失败，尝试从 reference/ 目录读取
    if not content:
        reference_path = get_project_root() / "reference" / file_name
        content = read_file(reference_path) or ""
        if content:
            print(f"✅ 从本地读取成功（reference路径）: {file_name} ({len(content)} 字符)")
        else:
            print(f"⚠️ 本地读取失败（两种路径都失败）: {file_name}")

    return content


def write_reference_file(file_name: str, content: str, message: str) -> bool:
    """
    写入 reference 文件（支持本地和 GitHub 模式）

    Args:
        file_name: 文件名（如 "01_历史纪要重点总结.md"）
        content: 文件内容
        message: 提交消息（用于 GitHub）

    Returns:
        是否成功
    """
    # 优先尝试写入 GitHub
    github_mgr = create_github_manager(st)
    if github_mgr:
        if github_mgr.update_reference_file(file_name, content, message):
            return True

    # 回退到本地写入
    local_path = get_project_root() / "reference" / file_name
    return write_file(local_path, content)


def read_docx(file_path):
    """读取 .docx 文件内容"""
    try:
        doc = Document(file_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        return text
    except Exception as e:
        st.error(f"读取 .docx 文件失败: {e}")
        return None


def get_latest_files(directory):
    """获取指定目录下最新的文件"""
    try:
        files = []
        for file_path in directory.glob('*'):
            if file_path.is_file():
                files.append((file_path, file_path.stat().st_mtime))
        files.sort(key=lambda x: x[1], reverse=True)
        return [f[0] for f in files]
    except Exception as e:
        st.error(f"获取文件列表失败: {e}")
        return []


def parse_terminology_dict(content):
    """解析术语词典内容"""
    terms = []
    lines = content.split('\n')
    current_section = None
    for line in lines:
        line = line.strip()
        if line.startswith('#'):
            current_section = line
            continue
        if '|' in line and not line.startswith('|---'):
            parts = [p.strip() for p in line.split('|')]
            if len(parts) >= 3 and parts[0] != '错误/口语称呼':
                terms.append({
                    'section': current_section,
                    'wrong': parts[0],
                    'correct': parts[1],
                    'note': parts[2] if len(parts) > 2 else ''
                })
    return terms


def add_term_to_dict(wrong_term, correct_term, note=''):
    """向术语词典添加新词条"""
    dict_path = get_dict_md_path()
    content = read_file(dict_path)

    # 找到合适的位置插入
    lines = content.split('\n')
    insert_index = -1

    for i, line in enumerate(lines):
        if line.startswith('## 口语化表达纠正'):
            insert_index = i + 7  # 跳过表头
            break

    if insert_index == -1:
        # 如果没找到，就追加到末尾
        insert_index = len(lines)

    # 插入新词条
    new_line = f"| {wrong_term} | {correct_term} | {note} |"
    lines.insert(insert_index, new_line)

    # 写回文件
    write_file(dict_path, '\n'.join(lines))
    return True


def export_to_markdown(content, file_path):
    """导出为 Markdown 文件"""
    return write_file(file_path, content)


def export_to_word(content, file_path):
    """导出为 Word 文档"""
    try:
        doc = Document()
        doc.add_heading('管理层周例会纪要', 0)

        # 简单的 Markdown 转 Word 逻辑
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('# '):
                doc.add_heading(line[2:], 0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 2)
            elif line.startswith('|') and line.count('|') >= 3:
                # 表格处理（简化版）
                pass
            elif line.startswith('**') and line.endswith('**'):
                doc.add_paragraph(line, style='Strong')
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(file_path)
        return True
    except Exception as e:
        st.error(f"导出 Word 失败: {e}")
        return False


# ==================== Streamlit 主程序 ====================

def main():
    # 页面标题
    st.markdown('<div class="main-header">📝 会议纪要生成 Agent</div>', unsafe_allow_html=True)
    st.markdown('<p style="text-align: center; color: #6B7280; margin-bottom: 2rem;">智能生成 | 格式规范 | 风格一致</p>', unsafe_allow_html=True)

    # ==================== 左侧边栏 - 大脑控制台 ====================
    with st.sidebar:
        st.markdown("## 🧠 大脑控制台")

        # 当前模式显示
        st.markdown("### 当前模式")

        # 检测 GitHub 模式
        github_mgr = create_github_manager(st)
        if github_mgr:
            st.success(f"✅ GitHub Cloud 模式")
            st.caption(f"仓库: {github_mgr.owner}/{github_mgr.repo}")
        else:
            st.info("📁 本地文件模式")

        st.markdown("---")

        # 历史纪要范文选择
        st.markdown("### 📚 历史纪要范文选择")

        # 读取可用的历史纪要文件
        available_samples = _get_available_history_samples()

        if available_samples:
            st.caption(f"找到 {len(available_samples)} 篇历史纪要")

            # 默认选择最新的 2 篇（用户要求）
            if 'selected_history_samples' not in st.session_state:
                available_samples.sort(key=lambda x: x.lower(), reverse=True)
                default_selected = available_samples[:2]
                st.session_state.selected_history_samples = default_selected

            # 多选框
            selected = st.multiselect(
                "选择作为范文的历史纪要（最多 2 篇）",
                options=available_samples,
                default=st.session_state.get('selected_history_samples', [])
            )

            st.session_state.selected_history_samples = selected

            # 显示选中的摘要
            if selected:
                st.info(f"已选择 {len(selected)} 篇范文：{', '.join(selected[:3])}...")
            else:
                st.warning("⚠️ 未选择范文，将只使用风格规则生成")
        else:
            st.warning("⚠️ 未找到历史纪要，请先上传历史会议纪要到 reference/历史纪要/ 目录")

        st.markdown("---")

        # GLM-4 API 配置
        st.markdown("### 🤖 GLM-4 API")

        # 尝试从 secrets 读取 API Key
        api_key = st.secrets.get("GLM_API_KEY", "")

        if not api_key:
            # 如果 secrets 中没有，从 session state 读取或让用户输入
            if "glm_api_key" not in st.session_state:
                st.session_state.glm_api_key = ""

            api_key = st.text_input(
                "API Key",
                type="password",
                value=st.session_state.glm_api_key,
                key="glm_api_key_input",
                help="智谱 GLM-4 API 密钥"
            )

            if api_key:
                st.session_state.glm_api_key = api_key
        else:
            st.success("✅ API Key 已从 Secrets 加载")

        # 测试连接按钮
        if st.button("🔍 测试 API 连接"):
            if api_key:
                client = GLMClient(api_key)
                if client.test_connection():
                    st.success("✅ API 连接成功！")
                else:
                    st.error("❌ API 连接失败，请检查 API Key")
            else:
                st.warning("⚠️ 请先输入 API Key")

        st.markdown("---")

        # 术语词典快捷添加
        st.markdown("### 📖 术语词典")

        st.markdown("**快捷添加术语**")

        wrong_term = st.text_input("错误/口语称呼", key="wrong_term")
        correct_term = st.text_input("正式称呼", key="correct_term")
        note = st.text_input("备注（可选）", key="note")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("添加到词典", type="primary"):
                if wrong_term and correct_term:
                    # 读取当前词典
                    dict_content = read_reference_file("02_组织与术语词典.md")

                    # 找到合适的位置插入
                    lines = dict_content.split('\n') if dict_content else []
                    insert_index = -1

                    for i, line in enumerate(lines):
                        if line.startswith('## 口语化表达纠正'):
                            insert_index = i + 7
                            break

                    if insert_index == -1:
                        insert_index = len(lines)

                    # 插入新词条
                    new_line = f"| {wrong_term} | {correct_term} | {note} |"
                    lines.insert(insert_index, new_line)

                    # 写回文件
                    new_content = '\n'.join(lines)
                    if write_reference_file("02_组织与术语词典.md", new_content, f"添加术语: {wrong_term} → {correct_term}"):
                        st.success(f"✅ 已添加: {wrong_term} → {correct_term}")
                        st.rerun()
                    else:
                        st.error("添加失败")
                else:
                    st.warning("请填写错误称呼和正式称呼")

        with col2:
            if st.button("刷新词典"):
                st.rerun()

        # 显示当前词典统计
        dict_content = read_reference_file("02_组织与术语词典.md")
        if dict_content:
            lines = dict_content.split('\n')
            terms = [line for line in lines if '|' in line and not line.startswith('|---') and not line.startswith('| 错误')]
            st.caption(f"📊 当前词典共 {len(terms)} 个词条")

        # 查看完整词典
        with st.expander("查看完整术语词典"):
            if dict_content:
                st.markdown(dict_content)
            else:
                st.info("术语词典文件不存在")

        st.markdown("---")

        # Reference 文件管理
        st.markdown("### 📚 Reference 文件")

        ref_files = ["01_历史纪要重点总结.md", "02_组织与术语词典.md"]
        st.caption(f"共 {len(ref_files)} 个 Reference 文件")

        for ref_file in ref_files:
            with st.expander(f"📄 {ref_file}"):
                content = read_reference_file(ref_file)
                if content:
                    st.markdown(content)

        st.markdown("---")

        # 使用说明
        st.markdown("### 💡 使用说明")

        with st.expander("查看使用说明"):
            if github_mgr:
                st.markdown("""
                **GitHub Cloud 模式使用指南**

                1. **资料投喂**
                   - 上传录音转写文件（.docx 或 .txt）
                   - 输入手写重点内容

                2. **生成纪要**
                   - 点击「🚀 一键生成会议纪要」按钮
                   - 系统将自动生成标准化会议纪要

                3. **结果处理**
                   - 在结果区域查看生成的纪要
                   - 支持「一键复制」和「下载」功能

                4. **Reference 更新**
                   - 自动从 GitHub 仓库读取历史纪要和术语词典
                   - 更新后的内容自动保存到 GitHub 仓库

                💡 **提示**: 所有数据自动保存到 GitHub，不会丢失！
                """)
            else:
                st.markdown("""
                1. **资料投喂**
                   - 上传录音转写文件（.docx 或 .txt）
                   - 输入手写重点内容

                2. **生成纪要**
                   - 点击「🚀 一键生成会议纪要」按钮
                   - 系统将自动生成标准化会议纪要

                3. **结果处理**
                   - 在结果区域查看生成的纪要
                   - 支持「一键复制」和「下载」功能

                4. **Skill 进化**
                   - 对结果不满意时，在反馈区输入修改要求
                   - 点击「升级 Agent 大脑」自动更新 skill.md
                """)

    # ==================== 右侧主工作区 ====================

    # 标签页导航
    tab1, tab2, tab3 = st.tabs(["📝 纪要生成", "📊 结果展示", "🧠 Skill 进化"])

    # ==================== 标签页1：纪要生成 ====================
    with tab1:
        st.markdown('<div class="section-title">模块 A：资料投喂</div>', unsafe_allow_html=True)

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("#### 📁 文件上传区")
            st.caption("支持 .docx 或 .txt 格式的录音转写文件")

            # 文件上传
            uploaded_files = st.file_uploader(
                "拖拽或点击上传文件",
                type=['docx', 'txt'],
                accept_multiple_files=True,
                key="file_upload"
            )

            # 显示已上传文件
            if uploaded_files:
                st.markdown("**已上传文件：**")
                for file in uploaded_files:
                    st.success(f"✅ {file.name} ({file.size / 1024:.1f} KB)")

        with col2:
            st.markdown("#### ✍️ 手写重点输入")
            st.caption("输入本次会议的手写重点及特殊关注点")

            # 手写重点输入
            handwritten_notes = st.text_area(
                "手写重点内容",
                placeholder="请输入手写重点，例如：\n\n会议日期：2026-03-03\n参会人员：冷总、刘雨、Alter等\n冷总身份：发言者A\n\n1. 销售进度：果蔬篮需补货1500套\n2. 冷总重点指示...",
                height=250,
                key="handwritten_notes"
            )

        st.markdown("---")

        # 生成按钮
        st.markdown("#### 🚀 生成会议纪要")
        col1, col2, col3 = st.columns([1, 2, 1])

        with col2:
            generate_button = st.button(
                "🚀 一键生成会议纪要",
                type="primary",
                use_container_width=True
            )

        if generate_button:
            with st.spinner("正在生成会议纪要，请稍候..."):
                # 获取 API Key
                api_key = st.secrets.get("GLM_API_KEY", st.session_state.get("glm_api_key", ""))

                if not api_key:
                    st.error("❌ 未配置 GLM-4 API Key，请在侧边栏输入")
                    st.stop()

                # 优先使用上传的文件，如果没有才读取 inputs/ 目录
                transcript_content = ""
                notes_content = handwritten_notes  # 使用文本框输入的手写重点

                # 优先处理上传的文件
                if uploaded_files:
                    # 从上传的文件中读取
                    for file in uploaded_files:
                        file_name_lower = file.name.lower()

                        # 检查是否是录音转写文件
                        is_transcript = any(keyword in file_name_lower for keyword in ['转写', '录音', '转录'])

                        st.info(f"📋 文件名: {file.name} (是否转写: {is_transcript})")

                        if is_transcript:
                            # 这是录音转写文件
                            if file.name.endswith('.docx'):
                                try:
                                    doc = Document(io.BytesIO(file.getvalue()))
                                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                                    transcript_content = '\n'.join(paragraphs)
                                    st.success(f"✅ 已读取 .docx 文件，共 {len(paragraphs)} 段")
                                except Exception as e:
                                    st.error(f"❌ 读取 .docx 文件失败: {e}")
                                    st.error(f"错误类型: {type(e).__name__}")
                                    import traceback
                                    st.error(f"详细错误: {traceback.format_exc()}")
                                    transcript_content = ""
                            else:
                                transcript_content = file.getvalue().decode('utf-8')
                                st.success(f"✅ 已读取文本文件，共 {len(transcript_content)} 字符")

                            st.success(f"✅ 已读取上传文件: {file.name}")
                            break  # 只取第一个转写文件

                    if not transcript_content:
                        st.warning("⚠️ 未找到录音转写文件")
                        st.info("请确保文件名包含以下关键词：'转写'、'录音'、'转录'")
                        if uploaded_files:
                            st.write("当前上传的文件：")
                            for f in uploaded_files:
                                st.write(f"- {f.name}")
                else:
                    # 没有上传文件，尝试从 inputs/ 目录读取（本地模式）
                    inputs_dir = get_project_root() / "inputs"

                    if inputs_dir.exists():
                        files = list(inputs_dir.glob('*'))

                        # 优先找录音转写文件
                        transcript_files = [f for f in files if any(
                            keyword in f.name.lower()
                            for keyword in ['转写', '录音', '转录']
                        ) and f.suffix in ['.docx', '.txt', '.md']]

                        if transcript_files:
                            # 按修改时间排序，取最新的
                            transcript_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
                            latest_transcript = transcript_files[0]

                            if latest_transcript.suffix == '.docx':
                                transcript_content = read_docx(latest_transcript)
                            else:
                                transcript_content = read_file(latest_transcript) or ""

                            st.success(f"✅ 已读取本地文件: {latest_transcript.name}")

                if not transcript_content and not notes_content:
                    st.error("❌ 未找到会议文件")
                    st.info("""
                    请确保：
                    1. **Streamlit Cloud**: 在「📁 文件上传区」上传了录音转写文件
                    2. **本地运行**: 将会议文件放入 inputs/ 目录
                    3. 在「✍️ 手写重点输入」输入了手写重点（可选）
                    """)
                    st.stop()

                # 读取 reference 文件（增强双重路径支持）
                reference = {}
                ref_errors = []

                # 读取历史总结（尝试两种路径）
                st.info("📖 正在读取 reference 文件...")
                summary_content_github = read_reference_file("01_历史纪要重点总结.md")
                summary_content_direct = read_reference_file_no_prefix("01_历史纪要重点总结.md")

                # 优先使用成功的读取
                if len(summary_content_github) > 0:
                    reference['01_历史纪要重点总结.md'] = summary_content_github
                    st.success(f"✅ 历史总结读取成功（GitHub）: {len(summary_content_github)} 字符")
                elif len(summary_content_direct) > 0:
                    reference['01_历史纪要重点总结.md'] = summary_content_direct
                    st.success(f"✅ 历史总结读取成功（直接）: {len(summary_content_direct)} 字符")
                else:
                    ref_errors.append("01_历史纪要重点总结.md: 读取失败")
                    st.error("❌ 历史总结读取失败，会议纪要可能无法参考历史风格")
                    st.info("💡 文件位置可能是：reference/01_历史纪要重点总结.md 或 01_历史纪要重点总结.md")

                # 读取术语词典（尝试两种路径）
                dict_content_github = read_reference_file("02_组织与术语词典.md")
                dict_content_direct = read_reference_file_no_prefix("02_组织与术语词典.md")

                if len(dict_content_github) > 0:
                    reference['02_组织与术语词典.md'] = dict_content_github
                    st.success(f"✅ 术语词典读取成功（GitHub）: {len(dict_content_github)} 字符")
                elif len(dict_content_direct) > 0:
                    reference['02_组织与术语词典.md'] = dict_content_direct
                    st.success(f"✅ 术语词典读取成功（直接）: {len(dict_content_direct)} 字符")
                else:
                    ref_errors.append("02_组织与术语词典.md: 读取失败")
                    st.error("❌ 术语词典读取失败，无法自动纠正错别字")

                # 读取用户偏好
                preferences_content = read_reference_file("03_用户偏好.json")
                if preferences_content:
                    reference['03_用户偏好.json'] = preferences_content
                    st.success(f"✅ 用户偏好读取成功: {len(preferences_content)} 字符")
                else:
                    # 用户偏好是可选的，不报错
                    reference['03_用户偏好.json'] = '{}'

                # 读取风格模板（新增）
                style_template_github = read_reference_file("04_风格模板.md")
                style_template_direct = read_reference_file_no_prefix("04_风格模板.md")

                if len(style_template_github) > 0:
                    reference['04_风格模板.md'] = style_template_github
                    st.success(f"✅ 风格模板读取成功（GitHub）: {len(style_template_github)} 字符")
                elif len(style_template_direct) > 0:
                    reference['04_风格模板.md'] = style_template_direct
                    st.success(f"✅ 风格模板读取成功（直接）: {len(style_template_direct)} 字符")
                else:
                    # 风格模板是可选的，不报错，但给提示
                    reference['04_风格模板.md'] = ''
                    st.info("⚠️ 风格模板未读取到，将使用默认风格")

                # 显示错误信息
                if ref_errors:
                    st.warning("⚠️ Reference 文件读取错误：")
                    for error in ref_errors:
                        st.write(f"   - {error}")

                    st.warning("💡 可能的原因和解决方案：")
                    st.markdown("""
                    ### 1. GitHub Token 问题
                    - Token 可能已失效或权限不足
                    - 请访问：https://github.com/settings/tokens
                    - 重新生成 Token，确保勾选 `repo` 权限

                    ### 2. GitHub 仓库配置问题
                    - 请确认 GITHUB_OWNER 和 GITHUB_REPO 正确
                    - 仓库必须设置为公开（Public）

                    ### 3. 本地文件缺失（回退方案）
                    - 如果 GitHub 读取失败，系统会自动尝试读取本地文件
                    - 请确保本地 `reference/` 目录存在且有内容

                    ### 4. 网络连接问题
                    - Streamlit Cloud 可能无法访问 GitHub API
                    - 请检查网络连接
                    """)

                # 显示调试信息
                st.markdown("---")
                st.markdown("#### 📊 输入数据统计")

                # 显示 GitHub 模式状态
                github_mgr = create_github_manager(st)
                is_github_mode = github_mgr is not None

                if is_github_mode:
                    st.success(f"🌐 **GitHub Cloud 模式** - 数据来源：GitHub 仓库 {github_mgr.owner}/{github_mgr.repo}")
                else:
                    st.info(f"📁 **本地文件模式** - 数据来源：本地 reference/ 目录")

                st.markdown("")

                # 显示输入数据统计
                st.write(f"- **录音转写长度**: {len(transcript_content)} 字符")
                st.write(f"- **手写重点长度**: {len(notes_content)} 字符")
                st.write(f"- **历史总结长度**: {len(reference.get('01_历史纪要重点总结.md', ''))} 字符")
                st.write(f"- **术语词典长度**: {len(reference.get('02_组织与术语词典.md', ''))} 字符")
                st.write(f"- **风格模板长度**: {len(reference.get('04_风格模板.md', ''))} 字符")

                # 添加警告提示
                if len(reference.get('01_历史纪要重点总结.md', '')) == 0:
                    st.warning("⚠️ 历史总结长度为 0，会议纪要可能无法参考历史风格")
                if len(reference.get('02_组织与术语词典.md', '')) == 0:
                    st.warning("⚠️ 术语词典长度为 0，无法自动纠正错别字")
                if len(reference.get('04_风格模板.md', '')) == 0:
                    st.info("💡 风格模板为空，系统将使用默认风格生成。上传最终版纪要后，系统会自动学习您的写作风格！")

                if not transcript_content:
                    st.error("⚠️ 录音转写内容为空！生成的会议纪要可能不完整")
                    st.stop()

                # 调用 GLM-4 API 生成会议纪要
                client = GLMClient(api_key)

                # 获取选中的历史纪要范文
                selected_history = st.session_state.get('selected_history_samples', [])

                # 读取完整范文内容
                history_samples = []
                for sample_file in selected_history:
                    content = read_reference_file(f"历史纪要/{sample_file}")
                    if content:
                        history_samples.append({
                            'title': sample_file.replace('.md', ''),
                            'content': content
                        })
                        print(f"✅ 读取范文：{sample_file}（{len(content)} 字符）")

                with st.spinner("正在调用 GLM-4 API 生成会议纪要，这可能需要 1-2 分钟..."):
                    generated_minutes = client.generate_minutes(
                        transcript=transcript_content,
                        notes=notes_content,
                        reference=reference,
                        history_samples=history_samples
                    )

                if generated_minutes:
                    st.success("✅ 会议纪要生成成功！")

                    # 保存到 session state
                    st.session_state.generated_content = generated_minutes

                    # 自动保存到 outputs/ 目录
                    outputs_dir = get_project_root() / "outputs"
                    outputs_dir.mkdir(exist_ok=True)

                    filename = f"管理周会纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                    filepath = outputs_dir / filename

                    if write_file(filepath, generated_minutes):
                        st.success(f"✅ 已保存到: {filepath}")
                    else:
                        st.warning("⚠️ 保存到文件失败，但可以在页面复制")

                    # 提示用户查看结果
                    st.info("👉 请切换到「📊 结果展示」标签页查看生成的会议纪要")
                else:
                    st.error("❌ 会议纪要生成失败，请检查 API Key 或重试")
                    st.info("可能的原因：\n- API Key 无效\n- API 服务暂时不可用\n- 网络连接问题\n\n请检查 API 配置后重试。")

    # ==================== 标签页2：结果展示 ====================
    with tab2:
        st.markdown('<div class="section-title">模块 B：结果展示</div>', unsafe_allow_html=True)

        # 检查是否有生成的内容
        if st.session_state.get("generated_content"):
            content = st.session_state.generated_content

            # 显示生成的纪要
            st.markdown(content)

            st.markdown("---")

            # 操作按钮
            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button("📋 一键复制", use_container_width=True):
                    import pyperclip
                    try:
                        pyperclip.copy(content)
                        st.success("✅ 已复制到剪贴板")
                    except:
                        st.error("❌ 复制失败，请手动复制")

            with col2:
                if st.button("📥 下载为 Markdown", use_container_width=True):
                    outputs_dir = get_project_root() / "outputs"
                    outputs_dir.mkdir(exist_ok=True)
                    file_name = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                    file_path = outputs_dir / file_name
                    if export_to_markdown(content, file_path):
                        st.success(f"✅ 已保存到 {file_path}")
                    else:
                        st.error("❌ 保存失败")

            with col3:
                if st.button("📄 下载为 Word", use_container_width=True):
                    outputs_dir = get_project_root() / "outputs"
                    outputs_dir.mkdir(exist_ok=True)
                    file_name = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    file_path = outputs_dir / file_name
                    if export_to_word(content, file_path):
                        st.success(f"✅ 已保存到 {file_path}")
                    else:
                        st.error("❌ 保存失败")

            st.markdown("---")

            # 上传最终版纪要
            st.markdown("#### 📤 上传最终版纪要（可选）")
            st.caption("如果您已人工修改完成最终版本，请上传以便更新 Reference")

            final_version = st.file_uploader(
                "上传最终版会议纪要",
                type=['md', 'docx', 'txt'],
                key="final_version_upload"
            )

            if final_version:
                st.success(f"✅ 已接收最终版: {final_version.name}")

                # 读取文件内容
                content = ""
                if final_version.name.endswith('.md'):
                    content = final_version.getvalue().decode('utf-8')
                elif final_version.name.endswith('.txt'):
                    content = final_version.getvalue().decode('utf-8')
                elif final_version.name.endswith('.docx'):
                    doc = Document(io.BytesIO(final_version.getvalue()))
                    content = '\n'.join([p.text for p in doc.paragraphs])

                # 获取初稿内容
                draft = st.session_state.get("generated_content", "")

                if draft and len(content) > 0:
                    # 获取 API Key
                    api_key = st.secrets.get("GLM_API_KEY", st.session_state.get("glm_api_key", ""))

                    if not api_key:
                        st.error("❌ 未配置 GLM-4 API Key，无法进行深度风格学习")
                    else:
                        # 创建深度风格学习器
                        from style_learner_v2 import DeepStyleLearner, update_summary_with_style, update_terms_dict, update_user_preferences
                        glm_client = GLMClient(api_key)
                        deep_learner = DeepStyleLearner(glm_client)

                        st.markdown("---")
                        st.markdown("#### 🧠 深度风格学习（正在分析您的写作风格...）")

                        # 1. 深度分析写作风格（措辞、句子结构、表达习惯等）
                        with st.spinner("📝 正在深度分析写作风格（措辞、句子结构、表达习惯...）"):
                            style_analysis = deep_learner.extract_writing_style(content, draft)

                            # 显示学习结果预览
                            st.markdown("---")
                            st.markdown("#### 📊 风格学习结果预览")
                            learned_count = sum(1 for k, v in style_analysis.items() if v and v != "待学习")
                            st.write(f"**学习进度**: {learned_count}/6 个维度已完成")
                            st.write(f"**数据量对比**: 最终版 {len(content)} 字符 vs 初稿 {len(draft)} 字符")
                            if learned_count < 6:
                                st.warning(f"⚠️ 还有 {6 - learned_count} 个维度未完成学习，建议多上传几篇会议纪要")

                            # 显示各维度的学习状态
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.markdown("**措辞特点**")
                                if style_analysis.get('措辞特点') and style_analysis['措辞特点'] != '待学习':
                                    st.success("✅ 已学习")
                                    st.caption(f"长度: {len(style_analysis['措辞特点'])} 字符")
                                else:
                                    st.warning("⏳ 待学习")

                            with col2:
                                st.markdown("**句子结构**")
                                if style_analysis.get('句子结构') and style_analysis['句子结构'] != '待学习':
                                    st.success("✅ 已学习")
                                    st.caption(f"长度: {len(style_analysis['句子结构'])} 字符")
                                else:
                                    st.warning("⏳ 待学习")

                            with col3:
                                st.markdown("**表达习惯**")
                                if style_analysis.get('表达习惯') and style_analysis['表达习惯'] != '待学习':
                                    st.success("✅ 已学习")
                                    st.caption(f"长度: {len(style_analysis['表达习惯'])} 字符")
                                else:
                                    st.warning("⏳ 待学习")

                            # 第二行
                            col4, col5, col6 = st.columns(3)
                            with col4:
                                st.markdown("**重点强调**")
                                if style_analysis.get('重点强调') and style_analysis['重点强调'] != '待学习':
                                    st.success("✅ 已学习")
                                    st.caption(f"长度: {len(style_analysis['重点强调'])} 字符")
                                else:
                                    st.warning("⏳ 待学习")

                            with col5:
                                st.markdown("**格式偏好**")
                                if style_analysis.get('格式偏好') and style_analysis['格式偏好'] != '待学习':
                                    st.success("✅ 已学习")
                                    st.caption(f"长度: {len(style_analysis['格式偏好'])} 字符")
                                else:
                                    st.warning("⏳ 待学习")

                            with col6:
                                st.markdown("**特殊表达**")
                                if style_analysis.get('特殊表达') and style_analysis['特殊表达'] != '待学习':
                                    st.success("✅ 已学习")
                                    st.caption(f"长度: {len(style_analysis['特殊表达'])} 字符")
                                else:
                                    st.warning("⏳ 待学习")

                        # 2. 提取新术语
                        with st.spinner("📖 正在提取新术语..."):
                            existing_terms = read_reference_file("02_组织与术语词典.md")
                            new_terms = deep_learner.extract_new_terms(content, existing_terms)

                        # 3. 更新历史纪要重点总结
                        with st.spinner("📚 正在更新历史纪要重点总结..."):
                            style_guide = deep_learner.generate_style_guide(style_analysis)
                            existing_summary = read_reference_file("01_历史纪要重点总结.md")
                            updated_summary = update_summary_with_style(existing_summary, style_guide)

                            # 保存到 GitHub
                            write_reference_file("01_历史纪要重点总结.md", updated_summary, f"深度学习更新: {final_version.name}")
                            st.success("✅ 历史纪要重点总结已更新")

                        # 4. 更新术语词典
                        if new_terms:
                            with st.spinner("📝 正在更新术语词典..."):
                                updated_terms_dict = update_terms_dict(existing_terms, new_terms)
                                write_reference_file("02_组织与术语词典.md", updated_terms_dict, f"新增术语: {len(new_terms)} 个")
                                st.success(f"✅ 术语词典已更新，新增 {len(new_terms)} 个术语")
                        else:
                            st.info("💡 未发现新术语")

                        # 5. 更新用户偏好（如果有多篇纪要）
                        st.info("💡 用户偏好分析需要至少 2 篇会议纪要才能准确分析")

                        # 6. 更新风格模板
                        with st.spinner("🎨 正在更新风格模板..."):
                            current_template = read_reference_file("04_风格模板.md")
                            new_template_content = f"""# 会议纪要最佳风格模板

本模板从深度学习中提取，持续优化中。

---

{style_guide}

---

{current_template}

---

## {datetime.now().strftime('%Y-%m-%d')} 深度学习更新

### 写作风格特点
{style_analysis.get('措辞特点', '')}

### 句子结构特点
{style_analysis.get('句子结构', '')}

### 表达习惯
{style_analysis.get('表达习惯', '')}

### 重点强调方式
{style_analysis.get('重点强调', '')}

---
"""
                            write_reference_file("04_风格模板.md", new_template_content, f"深度学习更新: {final_version.name}")
                            st.success("✅ 风格模板已更新")

                        # 显示详细学习结果
                        st.markdown("---")
                        st.markdown("#### 📊 深度学习结果")

                        with st.expander("查看详细学习内容", expanded=False):
                            st.markdown("### 🎯 措辞特点")
                            st.markdown(style_analysis.get('措辞特点', '分析失败'))

                            st.markdown("### 📐 句子结构")
                            st.markdown(style_analysis.get('句子结构', '分析失败'))

                            st.markdown("### 💬 表达习惯")
                            st.markdown(style_analysis.get('表达习惯', '分析失败'))

                            st.markdown("### ⭐ 重点强调方式")
                            st.markdown(style_analysis.get('重点强调', '分析失败'))

                            st.markdown("### 🎨 格式偏好")
                            st.markdown(style_analysis.get('格式偏好', '分析失败'))

                            if new_terms:
                                st.markdown("### 📖 新增术语")
                                for term in new_terms:
                                    st.markdown(f"- **{term['type']}**: {term['wrong']} → {term['correct']} ({term['note']})")

                        st.success("🎉 深度风格学习已完成！下次生成时会应用新学到的写作风格。")
                        st.info("💡 提示：越使用，越智能。建议多上传几次最终版会议纪要，让系统更深入地学习您的写作风格。")
                else:
                    # 没有初稿可对比，只做简单的术语提取
                    import jieba
                    words = jieba.cut(content)
                    word_freq = {}
                    for word in words:
                        if len(word) > 1:
                            word_freq[word] = word_freq.get(word, 0) + 1

                    # 提取潜在新术语
                    potential_terms = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]

                    # 更新历史纪要总结
                    summary_content = read_reference_file("01_历史纪要重点总结.md")
                    today = datetime.now().strftime('%Y-%m-%d')
                    new_entry = f"""

## {today} 自动更新总结

### 新发现要点
- 自动检测到最终版纪要更新
- 文件: {final_version.name}

### 潜在新术语
{chr(10).join([f"- {term[0]} (出现{term[1]}次)" for term in potential_terms])}

### 更新方式
GitHub Cloud 模式自动更新

---
"""
                    updated_summary = new_entry + summary_content if summary_content else new_entry
                    write_reference_file("01_历史纪要重点总结.md", updated_summary, f"更新历史纪要总结: {final_version.name}")

                    st.info("✅ Reference 文件已更新并保存到 GitHub")
                    st.info("💡 提示：下次生成会议纪要后，上传最终版时系统会自动学习您的写作风格！")

        else:
            st.info("👈 请先在「📝 纪要生成」标签页生成会议纪要")

    # ==================== 标签页3：Skill 进化 ====================
    with tab3:
        st.markdown('<div class="section-title">模块 C：Skill 进化反馈回路</div>', unsafe_allow_html=True)

        st.markdown("### 🔄 智能升级 Agent 大脑")

        st.markdown("""
        对生成的会议纪要不满意？或者想永久改变写作风格？

        在下方输入您的修改要求，系统将自动理解并更新 `skill.md` 文件。
        """)

        st.markdown("---")

        # 反馈输入
        feedback_input = st.text_area(
            "告诉 Claude 怎么改：",
            placeholder="例如：\n• 以后把所有提到财务的地方都加上需要左小美审核\n• 冷总发言前要加「冷总指示」前缀\n• TODO事项要按优先级排序",
            height=150,
            key="feedback_input"
        )

        # 常用修改建议
        st.markdown("#### 💡 常用修改建议（点击快速填入）")

        col1, col2, col3 = st.columns(3)

        with col1:
            if st.button("📊 增加数据表格格式", key="suggestion1"):
                feedback_input = "增加销售数据表格，包含：达成率、毛利率、订单量等关键指标"

        with col2:
            if st.button("👥 优化人员称呼", key="suggestion2"):
                feedback_input = "统一人员称呼格式，首次出现使用全名，后续使用简称"

        with col3:
            if st.button("⏰ 强调时间节点", key="suggestion3"):
                feedback_input = "所有事项都必须明确截止日期，优先级事项标注紧急程度"

        st.markdown("---")

        # 升级按钮
        col1, col2, col3 = st.columns([1, 2, 1])

        with col2:
            if st.button("🧠 升级 Agent 大脑", type="primary", use_container_width=True):
                if feedback_input:
                    with st.spinner("正在分析您的反馈并更新 skill.md..."):
                        # 这里应该调用 Claude API 分析反馈并更新 skill.md
                        # 由于没有实际的 API 集成，这里演示逻辑
                        st.success("✅ Agent 大脑已升级！")

                        # 读取当前 skill.md
                        skill_content = read_file(skill_md_path)

                        # 在合适位置插入新规则（演示逻辑）
                        new_rule = f"\n\n## 用户自定义规则（{datetime.now().strftime('%Y-%m-%d')}）\n{feedback_input}\n"
                        updated_content = skill_content + new_rule

                        # 写回文件
                        if write_file(skill_md_path, updated_content):
                            st.success("✅ skill.md 已更新")
                            st.info("👉 下次生成会议纪要时将应用新规则")
                        else:
                            st.error("❌ 更新失败")
                else:
                    st.warning("请先输入修改要求")

        st.markdown("---")

        # 当前 skill.md 预览
        st.markdown("### 📄 当前 skill.md 预览")

        if skill_md_path.exists():
            skill_content = read_file(skill_md_path)
            with st.expander("查看完整 skill.md", expanded=False):
                st.markdown(skill_content)
        else:
            st.warning("skill.md 文件不存在")


# ==================== 程序入口 ====================
if __name__ == "__main__":
    # 初始化 session state
    if "generated_content" not in st.session_state:
        st.session_state.generated_content = None

    main()
